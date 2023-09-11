import math
import os
import re

import fasttext
import numpy as np
import pandas as pd
from docx import Document

from utils import contain_string, csv_reader

train_path = 'C:\\Files\\Projects\\jupyter\\dummy.train'
test_path = 'C:\\Files\\Projects\\jupyter\\dummy.valid'
model_save_path = 'C:\\Files\\Projects\\jupyter\\models\\'
k_fold = 10

directory_paths = [
    'M:\\我的雲端硬碟\\UniMelb\\Research Project\\Open Coding\\20230522\\',  # 50
    'M:\\我的雲端硬碟\\UniMelb\\Research Project\\Open Coding\\20230601\\',  # 10
    'M:\\我的雲端硬碟\\UniMelb\\Research Project\\Open Coding\\20230606\\',  # 50
    'M:\\我的雲端硬碟\\UniMelb\\Research Project\\Open Coding\\20230607\\',  # 50
    'M:\\我的雲端硬碟\\UniMelb\\Research Project\\Open Coding\\20230608\\',  # 50
    'M:\\我的雲端硬碟\\UniMelb\\Research Project\\Open Coding\\20230609\\',  # 50
    'M:\\我的雲端硬碟\\UniMelb\\Research Project\\Open Coding\\20230610\\',  # 50
]

ignored_file_names = [
    'desktop.ini',
    'Completed'
]

colon_patterns = [
    ':',
    '：'
]

list_patterns = [
    '-',
    '*'
]

link_patterns = [
    'http://',
    'https://',
    'www.',
    '.com',
    '.net',
    '.org'
]

label_prefix = '__label__'

category_names = [
    'Introduction',
    'Threat model',
    'Scope',
    'Reporting procedure',
    'Handling procedure',
    'Secure communication',
    'Bug bounty program',
    'Known vulnerabilities',
    'Guideline',
    'Empty',
    'Subscription'
]


def get_fasttext_mappings():
    fasttext_mappings = dict()
    fasttext_mappings['__label__reporting-procedure'] = 'Reporting procedure'
    fasttext_mappings['__label__handling-procedure'] = 'Handling procedure'
    fasttext_mappings['__label__scope'] = 'Scope'
    fasttext_mappings['__label__secure-communication'] = 'Secure communication'
    fasttext_mappings['__label__threat-model'] = 'Threat model'
    fasttext_mappings['__label__empty'] = 'Empty'
    fasttext_mappings['__label__introduction'] = 'Introduction'
    fasttext_mappings['__label__bug-bounty-program'] = 'Bug bounty program'
    fasttext_mappings['__label__known-vulnerabilities'] = 'Known vulnerabilities'
    fasttext_mappings['__label__guideline'] = 'Guideline'
    fasttext_mappings['__label__subscription'] = 'Subscription'
    return fasttext_mappings


def get_lines(table, column_index):
    lines = None
    cells = table.columns[column_index].cells
    if len(cells) == 1:
        lines = cells[0].text
    return lines


def get_type_1_header(line, dummy, found):
    header = None
    if len(set([*line])) == 1 and ('-' in line[:4] or '=' in line[:4]):
        found = not found
        if found:
            header = ' '.join(dummy)
    return found, header


def get_type_2_header(line):
    header = None
    if line[:6] == '######':
        header = line[6:]
    elif line[:5] == '#####':
        header = line[5:]
    elif line[:4] == '####':
        header = line[4:]
    elif line[:3] == '###':
        header = line[3:]
    elif line[:2] == '##':
        header = line[2:]
    elif line[:1] == '#':
        header = line[1:]
    if header is not None:
        header = header.strip()
    return header


def preprocess(s):
    s = re.sub('[^\w\d ]', '', s)
    s = re.sub(' +', ' ', s)
    s = s.strip()
    s = s.lower()
    return s


def df_dummy(directory_path, file_name, file_headers, file_paragraphs, file_categories):
    # print(paragraph_categories)
    # if len(file_categories) != len(file_paragraphs):

    data = []

    # print(f'{directory_path}{file_name}')

    for i in range(len(file_categories)):
       # if len(file_headers[i]) > 0:
        #     lines.append(' '.join(labels) + ' ' + headers[i] + ' ' + paragraphs[i])
        # else:
        #     lines.append(' '.join(labels) + ' ' + paragraphs[i])


        paragraph = preprocess(file_paragraphs[i])
        labels = map(lambda x: x.lower().replace(' ', '-'), file_categories[i])
        labels = map(lambda x: f'{label_prefix}{x}', labels)
        labels = ' '.join(labels)

        # print()
        # print(f'FILE_NAME           : {directory_path}{file_name}')
        # print(f'HEADER              : {file_headers[i]}')
        # print(f'PARAGRAPH           : {file_paragraphs[i]}')
        # print(f'CATEGORIES          : {file_categories[i]}')

        data.append({
            'name': file_name,
            'headers': file_headers[i],
            'paragraph': paragraph,
            'labels': labels
        })

    return pd.DataFrame(data)


def get_headers_and_paragraphs(content, append_header_to_paragraph=False):
    headers = []
    paragraphs = []
    previous_header = ''
    header_lines = []
    paragraph_lines = []
    found = False
    for line in content.split('\n'):
        line = line.strip()
        found, header = get_type_1_header(line, header_lines, found)
        if header is not None:
            header_lines.clear()
            paragraph_lines.clear()
        else:
            header_lines.append(line)
            header = get_type_2_header(line)
        if len(line) == 0:
            found = False
            header_lines.clear()
        if header is None:
            if len(line) > 0:
                paragraph_lines.append(line)
            elif len(paragraph_lines) > 0:
                headers.append(previous_header)
                paragraph = '\n'.join(map(lambda x: x.strip(), paragraph_lines)).strip()
                # if append_header_to_paragraph:
                #     paragraph = f'{previous_header} {paragraph}'.strip()
                paragraphs.append(paragraph)
                paragraph_lines = []
        if header is not None:
            previous_header = header
    if len(paragraph_lines) > 0:
        headers.append(previous_header)
        paragraph = '\n'.join(map(lambda x: x.strip(), paragraph_lines)).strip()
        paragraphs.append(paragraph)
    return combine_headers_and_paragraphs(headers, paragraphs, append_header_to_paragraph)


def combine_headers_and_paragraphs(headers, paragraphs, append_header_to_paragraph):
    new_headers = []
    new_paragraphs = []
    header_buffer = []
    paragraph_buffer = []
    for i in range(len(paragraphs)):
        header = headers[i]
        paragraph = paragraphs[i]
        if len(header_buffer) == 0:
            if paragraph[-1] in colon_patterns:
                header_buffer.append(header)
                paragraph_buffer.append(paragraph)
            else:
                new_headers.append(header)
                new_paragraph = paragraph
                if append_header_to_paragraph:
                    new_paragraph = f'{header} {new_paragraph}'.strip()
                new_paragraphs.append(new_paragraph)
        else:
            first_word = paragraph.split(' ')[0]
            unordered_list_item = first_word in list_patterns
            ordered_list_item = first_word[-1] == '.' and first_word[:-1].isdigit()
            link = contain_string(first_word, link_patterns)
            if unordered_list_item or ordered_list_item or link:
                if len(header_buffer) > 0 and header != header_buffer[0]:
                    new_header = header_buffer[0]
                    new_headers.append(new_header)
                    new_paragraph = '\n'.join(paragraph_buffer)
                    if append_header_to_paragraph:
                        new_paragraph = f'{new_header} {new_paragraph}'.strip()
                    new_paragraphs.append(new_paragraph)
                    header_buffer.clear()
                    paragraph_buffer.clear()
                header_buffer.append(header)
                paragraph_buffer.append(paragraph)
            else:
                new_header = header_buffer[0]
                new_headers.append(new_header)
                new_paragraph = '\n'.join(paragraph_buffer)
                if append_header_to_paragraph:
                    new_paragraph = f'{new_header} {new_paragraph}'.strip()
                new_paragraphs.append(new_paragraph)
                header_buffer.clear()
                paragraph_buffer.clear()
                if paragraph[-1] in colon_patterns:
                    header_buffer.append(header)
                    paragraph_buffer.append(paragraph)
                else:
                    new_headers.append(header)
                    new_paragraph = paragraph
                    if append_header_to_paragraph:
                        new_paragraph = f'{header} {new_paragraph}'.strip()
                    new_paragraphs.append(new_paragraph)
    if len(header_buffer) > 0:
        new_header = header_buffer[0]
        new_headers.append(new_header)
        new_paragraph = '\n'.join(paragraph_buffer)
        if append_header_to_paragraph:
            new_paragraph = f'{new_header} {new_paragraph}'.strip()
        new_paragraphs.append(new_paragraph)
    return new_headers, new_paragraphs


def get_categories(content):
    categories = []
    category_lines = []
    for line in content.split('\n'):
        if len(line) > 0:
            category_lines.append(line)
        elif len(category_lines) > 0:
            categories.append(category_lines)
            category_lines = []
    if len(category_lines) > 0:
        categories.append(category_lines)
    return categories


def get_directory_paths():
    directory_paths = []
    parent_directory_path = 'M:\\我的雲端硬碟\\UniMelb\\Research Project\\Open Coding\\'
    for path in os.listdir(parent_directory_path):
        path = f'{parent_directory_path}{path}\\'
        if os.path.isdir(path) and 'security-related issues' not in path:
            directory_paths.append(path)
    return directory_paths


def get_docx_content(file_path):
    content = None
    file = open(file_path, 'rb')
    paragraphs = Document(file).paragraphs
    if len(paragraphs) == 1:
        content = paragraphs[0].text
    file.close()
    return content


def get_docx_file_tuple(file_path):
    file_tuple = None
    file = open(file_path, 'rb')
    tables = Document(file).tables
    if len(tables) == 1:
        table = tables[0]
        categories = get_categories(get_lines(table, 0))
        headers, paragraphs = get_headers_and_paragraphs(get_lines(table, 1), True)
        if len(categories) == len(paragraphs):
            file_tuple = (headers, paragraphs, categories)
    file.close()
    return file_tuple


def get_csv_file_tuple(file_path):
    file_tuple = None
    headers = []
    paragraphs = []
    categories = []
    # print(file_path)
    for header, paragraph, paragraph_categories in csv_reader(file_path):
        headers.append(header)
        paragraphs.append(paragraph)
        paragraph_categories = list(map(lambda x: x.strip(), paragraph_categories.split(',')))
        if len(paragraph_categories) > 0:
            categories.append(paragraph_categories)
    if len(categories) == len(paragraphs):
        file_tuple = (headers, paragraphs, categories)
    return file_tuple


def get_df_list():
    df_list = []
    for directory_path in get_directory_paths():
        file_names = os.listdir(directory_path)
        for file_name in file_names:
            if file_name in ignored_file_names:
                continue
            file_tuple = None
            file_path = f'{directory_path}{file_name}'
            file_extension = file_name.split('.')[-1]
            if file_extension == 'docx':
                file_tuple = get_docx_file_tuple(file_path)
            elif file_extension == 'csv':
                file_tuple = get_csv_file_tuple(file_path)
            if file_tuple is not None:
                headers, paragraphs, categories = file_tuple
                df_list.append(df_dummy(directory_path, file_name, headers, paragraphs, categories))
            else:
                print(f'INVALID FILE PATH: {file_path}')
    print(f'df_list SIZE: {len(df_list)}')
    return df_list


def get_dataset(random=True):
    dataset = pd.concat(get_df_list(), ignore_index=True)
    if random:
        dataset = dataset.sample(frac=1)
    return dataset


def get_segments(dataset):
    segments = []
    dataset_size = len(dataset)
    segment_size = math.ceil(dataset_size / k_fold)
    for i in range(k_fold):
        start = i * segment_size
        if dataset_size - start >= segment_size:
            end = start + segment_size
        else:
            end = dataset_size
        segments.append(dataset[start:end])
    return segments


def get_training_and_test_set(segments, k_fold):
    training_set = None
    test_set = None
    training_set_segments = []
    for i in range(len(segments)):
        if i == k_fold:
            test_set = segments[i]
        else:
            training_set_segments.append(segments[i])
    if len(training_set_segments) > 0:
        training_set = pd.concat(training_set_segments, ignore_index=True)
    return training_set, test_set


def get_fasttext_model(training_set, test_set, save_path):
    columns = ['labels', 'paragraph']
    fmt = '%s %s'
    encoding = 'utf-8'
    np.savetxt(train_path, training_set[columns].values, fmt=fmt, encoding=encoding)
    np.savetxt(test_path, test_set[columns].values, fmt=fmt, encoding=encoding)
    model = fasttext.train_supervised(input=train_path, lr=0.5, epoch=25, wordNgrams=2, bucket=200000, dim=300, loss='ova')
    model.save_model(save_path)
    return model


def get_distinct_categories(file_path):
    my_dict = dict()
    file = open(file_path, 'r')
    for line in file.readlines():
        line = line.strip()
        if len(line) > 0:
            for category in map(lambda x: x.strip(), line.split(',')):
                my_dict[category] = 1
    return my_dict.keys()


if __name__ == '__main__':
    categories = get_distinct_categories('C:\\Users\\WAN Tung Lok\\Desktop\\pham.txt')
    for category in categories:
        print(category)